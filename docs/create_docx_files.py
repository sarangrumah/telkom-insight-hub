"""
Generate 4 DOCX documents for Telkom Insight Hub:
1. TECHNICAL_SPECIFICATION.docx
2. BLUEPRINT.docx
3. BUSINESS_REQUIREMENTS.docx
4. MONTHLY_PROGRESS_REPORT.docx
All with sign-off boxes.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(__file__)

# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "E11A27")
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_bg(cell, "F5F5F5")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_signoff_box(doc, title="Lembar Persetujuan"):
    """Add a sign-off box section."""
    doc.add_page_break()
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xE1, 0x1A, 0x27)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Peran", "Nama", "Tanda Tangan", "Tanggal"]
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "E11A27")

    roles = [
        "Disusun oleh (Prepared by)",
        "Diperiksa oleh (Reviewed by)",
        "Disetujui oleh (Approved by - Teknis)",
        "Disetujui oleh (Approved by - Manajemen)",
    ]
    for i, role in enumerate(roles):
        table.rows[i + 1].cells[0].text = role
        table.rows[i + 1].cells[1].text = ""
        table.rows[i + 1].cells[2].text = ""
        table.rows[i + 1].cells[3].text = ""
        # Set minimum height for signature
        for cell in table.rows[i + 1].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            cell.height = Cm(2)

    doc.add_paragraph("")
    doc.add_paragraph("Catatan Revisi:", style="Heading 3")
    rev_table = doc.add_table(rows=4, cols=5)
    rev_table.style = 'Table Grid'
    rev_headers = ["Versi", "Tanggal", "Penulis", "Deskripsi Perubahan", "Status"]
    for i, h_text in enumerate(rev_headers):
        cell = rev_table.rows[0].cells[i]
        cell.text = h_text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "333333")
    rev_table.rows[1].cells[0].text = "1.0"
    rev_table.rows[1].cells[1].text = "April 2026"
    rev_table.rows[1].cells[2].text = "Tim Pengembangan"
    rev_table.rows[1].cells[3].text = "Dokumen awal"
    rev_table.rows[1].cells[4].text = "Draft"

def add_cover_page(doc, title, subtitle, doc_id):
    """Add a cover page."""
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KEMENTERIAN KOMUNIKASI DAN DIGITAL")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Direktorat Jenderal Penyelenggaraan Pos dan Informatika")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xE1, 0x1A, 0x27)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TELKOM INSIGHT HUB")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(3):
        doc.add_paragraph("")

    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Nomor Dokumen:", doc_id),
        ("Versi:", "1.0"),
        ("Tanggal:", "April 2026"),
        ("Klasifikasi:", "Internal — Terbatas"),
        ("Status:", "Draft untuk Persetujuan"),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        for cell in info_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    doc.add_page_break()

def section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xE1, 0x1A, 0x27) if level == 1 else RGBColor(0x1A, 0x1A, 0x2E)
    return h

# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: TECHNICAL_SPECIFICATION.docx
# ═══════════════════════════════════════════════════════════════════════════

def create_technical_specification():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_cover_page(doc, "TECHNICAL\nSPECIFICATION", "Spesifikasi Teknis Sistem Panel Manajemen Data Telekomunikasi", "TIH-TS-2026-001")

    # TOC placeholder
    section_heading(doc, "Daftar Isi")
    toc_items = [
        "1. Pendahuluan", "2. Arsitektur Sistem", "3. Technology Stack",
        "4. Frontend Architecture", "5. Backend Architecture", "6. Database Design",
        "7. API Specification", "8. Authentication & Authorization", "9. Security Implementation",
        "10. Integration Architecture", "11. Performance Requirements", "12. Deployment Architecture",
        "13. Testing Strategy", "14. Monitoring & Logging"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # 1. Pendahuluan
    section_heading(doc, "1. Pendahuluan")
    doc.add_paragraph(
        "Dokumen ini merupakan spesifikasi teknis lengkap untuk sistem Telkom Insight Hub, "
        "sebuah platform manajemen data telekomunikasi dan analitik strategis yang dikembangkan "
        "untuk Kementerian Komunikasi dan Digital Republik Indonesia. Dokumen ini mencakup "
        "arsitektur sistem, teknologi yang digunakan, desain database, spesifikasi API, "
        "implementasi keamanan, dan strategi deployment."
    )

    section_heading(doc, "1.1 Tujuan Dokumen", level=2)
    doc.add_paragraph(
        "Spesifikasi teknis ini bertujuan untuk:\n"
        "- Mendokumentasikan arsitektur dan desain teknis sistem\n"
        "- Menjadi acuan bagi tim pengembangan dalam implementasi\n"
        "- Menyediakan referensi teknis untuk maintenance dan enhancement\n"
        "- Memfasilitasi knowledge transfer kepada tim operasional"
    )

    section_heading(doc, "1.2 Ruang Lingkup", level=2)
    doc.add_paragraph(
        "Sistem Telkom Insight Hub mencakup pengelolaan data perizinan telekomunikasi "
        "untuk 8 jenis layanan (Jasa, Jaringan, Telsus, ISR, SKLO, LKO, Penomoran, Tarif), "
        "integrasi dengan BPS dan KOMINFO, analisis potensi pasar telekomunikasi, "
        "serta modul pendukung berupa registrasi perusahaan, support ticketing, dan FAQ management."
    )

    # 2. Arsitektur Sistem
    section_heading(doc, "2. Arsitektur Sistem")
    doc.add_paragraph(
        "Sistem menggunakan arsitektur 3-tier (Three-Tier Architecture) yang terdiri dari "
        "Client Layer (Frontend), Application Layer (Backend API), dan Data Layer (Database & External Services)."
    )

    section_heading(doc, "2.1 Arsitektur Overview", level=2)
    add_styled_table(doc,
        ["Layer", "Teknologi", "Fungsi", "Port"],
        [
            ["Client (Frontend)", "React 18 + TypeScript + Vite", "User Interface & Interaction", "5173 (dev) / 8080"],
            ["Application (Backend)", "Node.js + Express 4.19", "Business Logic & API", "4000"],
            ["Data (Database)", "PostgreSQL 15+", "Data Persistence", "5432"],
            ["Real-time", "WebSocket (ws 8.18)", "Live Updates & Notifications", "4000 (WS)"],
            ["External APIs", "BPS, KOMINFO, e-Telekomunikasi", "Third-party Data Sources", "Various"],
        ],
        col_widths=[1.5, 2.0, 2.0, 1.0]
    )

    section_heading(doc, "2.2 Communication Patterns", level=2)
    doc.add_paragraph(
        "- Frontend \u2194 Backend: RESTful HTTP API dengan JSON payload\n"
        "- Real-time Updates: WebSocket bidirectional connection\n"
        "- Backend \u2194 External APIs: HTTP/HTTPS dengan API key authentication\n"
        "- Background Jobs: node-cron scheduler untuk periodic sync tasks\n"
        "- Event System: Internal event emitter untuk cross-module communication"
    )

    # 3. Technology Stack
    section_heading(doc, "3. Technology Stack")

    section_heading(doc, "3.1 Frontend Dependencies", level=2)
    add_styled_table(doc,
        ["Package", "Version", "Kategori", "Fungsi"],
        [
            ["react", "18.3.1", "Core", "UI Component Library"],
            ["typescript", "5.8.3", "Core", "Type Safety"],
            ["vite", "5.4.19", "Build", "Development Server & Bundler"],
            ["react-router-dom", "6.30.1", "Routing", "Client-side Navigation"],
            ["@tanstack/react-query", "5.83.0", "State", "Server State Management & Caching"],
            ["tailwindcss", "3.4.17", "Styling", "Utility-first CSS Framework"],
            ["shadcn/ui (Radix)", "Various", "UI", "47+ Accessible UI Components"],
            ["recharts", "2.15.4", "Charts", "Data Visualization & Analytics"],
            ["mapbox-gl", "3.14.0", "Maps", "Interactive Map Visualization"],
            ["react-hook-form", "7.61.1", "Forms", "Form State & Validation"],
            ["zod", "3.25.76", "Validation", "Schema Validation"],
            ["axios", "1.13.2", "HTTP", "API Client"],
            ["lucide-react", "0.462.0", "Icons", "Icon Library"],
            ["date-fns", "3.6.0", "Utility", "Date Manipulation"],
            ["sonner", "1.7.4", "UI", "Toast Notifications"],
            ["xlsx", "0.18.5", "Export", "Excel Import/Export"],
        ],
        col_widths=[1.8, 0.8, 0.8, 3.0]
    )

    section_heading(doc, "3.2 Backend Dependencies", level=2)
    add_styled_table(doc,
        ["Package", "Version", "Kategori", "Fungsi"],
        [
            ["express", "4.19.2", "Core", "HTTP Server Framework"],
            ["pg", "8.13.1", "Database", "PostgreSQL Client"],
            ["jsonwebtoken", "9.0.2", "Auth", "JWT Token Generation & Verification"],
            ["bcryptjs", "2.4.3", "Auth", "Password Hashing"],
            ["helmet", "8.1.0", "Security", "HTTP Security Headers"],
            ["cors", "2.8.5", "Security", "Cross-Origin Resource Sharing"],
            ["express-rate-limit", "7.4.0", "Security", "Request Rate Limiting"],
            ["multer", "1.4.5", "Upload", "File Upload Handling"],
            ["node-cron", "4.2.1", "Jobs", "Background Job Scheduling"],
            ["ws", "8.18.0", "Real-time", "WebSocket Server"],
            ["cookie-parser", "1.4.7", "Middleware", "Cookie Parsing"],
            ["dotenv", "16.4.5", "Config", "Environment Variable Management"],
        ],
        col_widths=[1.8, 0.8, 0.8, 3.0]
    )

    # 4. Frontend Architecture
    section_heading(doc, "4. Frontend Architecture")

    section_heading(doc, "4.1 Routing Structure", level=2)
    add_styled_table(doc,
        ["Route", "Component", "Akses", "Deskripsi"],
        [
            ["/", "Index/Landing", "Public", "Homepage dengan overview sistem"],
            ["/register", "EnhancedPublicRegister", "Public", "Registrasi perusahaan baru"],
            ["/public-data", "PublicDataView", "Public", "Pencarian data perizinan"],
            ["/services/:type", "Service Pages (8)", "Public", "Detail per jenis layanan"],
            ["/faq", "FAQ", "Public", "Frequently Asked Questions"],
            ["/support", "Support", "Public", "Form dukungan teknis"],
            ["/dashboard", "EnhancedDashboard", "Protected", "Dashboard utama dengan KPI"],
            ["/data-management", "DataManagement", "Protected", "CRUD data telekomunikasi"],
            ["/bps-config", "BPSConfiguration", "Protected", "Konfigurasi API BPS"],
            ["/bps-visualization", "BPSDataVisualization", "Protected", "Visualisasi data statistik"],
            ["/bps-surveys", "BPSSurveyManagement", "Protected", "Manajemen variabel BPS"],
            ["/telecom-potential", "TelecomPotential", "Protected", "Analisis potensi pasar"],
            ["/integrations", "IntegrationsDashboard", "Protected", "Konfigurasi integrasi"],
            ["/users", "UserManagement", "Admin", "Manajemen pengguna"],
            ["/permissions", "PermissionManagement", "Admin", "Manajemen hak akses"],
            ["/company-management", "CompanyManagement", "Admin", "Verifikasi perusahaan"],
            ["/support-admin", "AdminTickets", "Admin", "Dashboard tiket admin"],
            ["/admin-faq", "AdminFAQ", "Admin", "Editor FAQ"],
        ],
        col_widths=[1.5, 1.8, 0.8, 2.5]
    )

    section_heading(doc, "4.2 Component Architecture", level=2)
    doc.add_paragraph(
        "Frontend menggunakan component-based architecture dengan 55+ reusable components. "
        "Komponen dibagi menjadi:\n"
        "- UI Primitives (47 shadcn/ui components): Button, Card, Dialog, Table, Form, dll.\n"
        "- Feature Components: BPSDataVisualization, TelecomPotentialDashboard, ExcelImportDialog\n"
        "- Layout Components: AppLayout (sidebar navigation), AuthPage\n"
        "- Guard Components: PermissionGuard (RBAC wrapper)\n"
        "- Hook-based Logic: useAuth, usePermissions, useTelekomData, useRealtimeTickets"
    )

    section_heading(doc, "4.3 State Management", level=2)
    doc.add_paragraph(
        "- Server State: TanStack React Query v5 untuk caching, refetching, optimistic updates\n"
        "- Auth State: React Context (AuthProvider) dengan JWT token management\n"
        "- Form State: React Hook Form + Zod schema validation\n"
        "- Real-time State: WebSocket client untuk live ticket updates\n"
        "- UI State: Local component state untuk modals, filters, pagination"
    )

    # 5. Backend Architecture
    section_heading(doc, "5. Backend Architecture")

    section_heading(doc, "5.1 Server Configuration", level=2)
    doc.add_paragraph(
        "Express server berjalan pada port 4000 dengan konfigurasi:\n"
        "- Helmet security headers (CSP, HSTS, X-Frame-Options)\n"
        "- CORS whitelist (configurable via CORS_ORIGIN env)\n"
        "- Rate limiting pada endpoint sensitif\n"
        "- Cookie parser untuk session management\n"
        "- Multer untuk file upload (storage/uploads/)\n"
        "- WebSocket server untuk real-time communication"
    )

    section_heading(doc, "5.2 Route Organization", level=2)
    add_styled_table(doc,
        ["Router File", "Base Path", "Endpoints", "Deskripsi"],
        [
            ["auth (index.js)", "/panel/api/auth/*", "6", "Authentication & registration"],
            ["telekom-data.js", "/panel/api/telekom-data/*", "5", "Telekom data CRUD"],
            ["bps.js", "/panel/api/bps/*", "12", "BPS statistical integration"],
            ["telecom-potential.js", "/panel/api/telecom-potential/*", "10", "Market potential scoring"],
            ["registration.js", "/panel/api/registration/*", "4", "Company registration"],
            ["adminVerification.js", "/panel/api/admin/*", "8", "Company verification"],
            ["integrations.js", "/panel/api/integrations/*", "5", "Sync configuration"],
            ["kominfo-sync.js", "/panel/api/kominfo-sync/*", "4", "KOMINFO tariff sync"],
            ["tickets.js", "Built-in", "5", "Support ticketing"],
            ["messages.js", "Built-in", "3", "Ticket messages"],
            ["tarif.js", "/panel/api/tarif/*", "2", "Tariff data"],
            ["sklo.js", "/panel/api/sklo/*", "2", "SKLO service data"],
            ["setup-db.js", "Initialization", "1", "Database setup"],
        ],
        col_widths=[1.5, 2.0, 0.8, 2.2]
    )

    section_heading(doc, "5.3 Service Layer", level=2)
    add_styled_table(doc,
        ["Service", "File", "Fungsi"],
        [
            ["BPS Data Service", "bpsDataService.js", "BPS data CRUD operations"],
            ["BPS API Fetcher", "bpsAPIFetcherService.js (23KB)", "Fetching data dari webapi.bps.go.id"],
            ["BPS Data Normalizer", "bpsDataNormalizerService.js (17KB)", "Normalisasi data BPS ke format internal"],
            ["KOMINFO Tarif Service", "kominfoTarifService.js (9KB)", "Client API tariftel.komdigi.go.id"],
            ["Tarif Sync Service", "tarifSyncService.js (15KB)", "Orchestration sinkronisasi tarif"],
            ["Telecom Potential V1", "telecomPotentialService.js (15KB)", "Scoring potensi pasar versi 1"],
            ["Telecom Potential V2", "telecomPotentialV2Service.js (19KB)", "Scoring dengan BPS data blending"],
            ["Background Job", "backgroundJobService.js (6KB)", "Cron job runner untuk scheduled tasks"],
            ["e-Telekomunikasi Client", "etelekom-client.js", "Integration adapter"],
            ["OSS Integration", "oss-integration.js", "OSS adapter (framework ready)"],
            ["POSTEL Integration", "postel-integration.js", "POSTEL legacy adapter"],
            ["SDPPI Integration", "sdppi-integration.js", "SDPPI spectrum adapter"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    # 6. Database Design
    section_heading(doc, "6. Database Design")

    section_heading(doc, "6.1 Schema Overview", level=2)
    doc.add_paragraph(
        "Database menggunakan PostgreSQL dengan 40+ tabel yang diorganisasi dalam beberapa domain:"
    )

    add_styled_table(doc,
        ["Domain", "Tabel", "Jumlah", "Deskripsi"],
        [
            ["Authentication", "auth.users, auth.sessions, auth.refresh_tokens, auth.identities", "4", "Autentikasi dan sesi"],
            ["User & Company", "profiles, user_roles, companies, company_documents, person_in_charge, pic_documents, user_profiles", "7", "Data pengguna dan perusahaan"],
            ["Telekom Data", "telekom_data, services, sub_services, license_services, provinces, kabupaten, indonesian_regions", "7", "Data perizinan telekomunikasi"],
            ["BPS Statistical", "bps_config, bps_variables, bps_monitored_areas, bps_statistical_data, bps_sync_history, bps_api_requests", "6", "Data statistik BPS"],
            ["Tariff & Sync", "kominfo_tarif_data, sync_status, sync_configurations", "3", "Data tarif dan sinkronisasi"],
            ["Market Potential", "telecom_potential_config, telecom_potential_area_scores", "2", "Scoring potensi pasar"],
            ["Support", "tickets, ticket_messages, ticket_assignments, ticket_sla_metrics, support_tickets", "5", "Sistem tiket"],
            ["Permissions", "permissions, permission_templates, record_permissions, modules, fields", "5", "Hak akses"],
            ["Audit", "audit_logs, activity_logs, api_integration_logs, login_attempts", "4", "Audit trail"],
            ["FAQ & Others", "faq_categories, faqs, captcha_sessions", "3", "FAQ dan lainnya"],
        ],
        col_widths=[1.2, 2.5, 0.5, 2.3]
    )

    section_heading(doc, "6.2 Key Indexes", level=2)
    doc.add_paragraph(
        "- Time-series: (area_type, area_id, variable_id, year) pada bps_statistical_data\n"
        "- Full-text search: company names, license numbers pada telekom_data\n"
        "- Status filtering: company status, ticket status\n"
        "- Activity tracking: timestamps pada audit_logs dan activity_logs\n"
        "- Unique constraints: email, NIB, NPWP pada companies"
    )

    # 7. API Specification
    section_heading(doc, "7. API Specification")
    doc.add_paragraph(
        "Seluruh API endpoint menggunakan base path /panel/api/ dengan format JSON. "
        "Authentication menggunakan JWT Bearer token pada header Authorization."
    )

    section_heading(doc, "7.1 Authentication Endpoints", level=2)
    add_styled_table(doc,
        ["Method", "Endpoint", "Auth", "Deskripsi"],
        [
            ["POST", "/panel/api/auth/register", "No", "Registrasi pengguna baru"],
            ["POST", "/panel/api/auth/login", "No", "Login dengan email/password"],
            ["POST", "/panel/api/auth/login-etelekomunikasi", "No", "Login via e-Telekomunikasi"],
            ["POST", "/panel/api/auth/refresh", "Cookie", "Refresh access token"],
            ["POST", "/panel/api/auth/logout", "Yes", "Logout dan clear session"],
            ["GET", "/panel/api/auth/check-email", "No", "Cek ketersediaan email"],
        ],
        col_widths=[0.7, 2.5, 0.6, 2.5]
    )

    section_heading(doc, "7.2 Telekom Data Endpoints", level=2)
    add_styled_table(doc,
        ["Method", "Endpoint", "Auth", "Deskripsi"],
        [
            ["GET", "/panel/api/telekom-data", "Yes", "List data dengan pagination & filter"],
            ["POST", "/panel/api/telekom-data", "Yes", "Create record baru"],
            ["PATCH", "/panel/api/telekom-data/:id", "Yes", "Update record"],
            ["DELETE", "/panel/api/telekom-data/:id", "Yes", "Delete record"],
            ["GET", "/panel/api/telekom-data/:id", "Yes", "Get single record detail"],
            ["POST", "/panel/api/telekom-data/excel", "Yes", "Bulk import dari Excel"],
        ],
        col_widths=[0.7, 2.5, 0.6, 2.5]
    )

    section_heading(doc, "7.3 BPS Integration Endpoints", level=2)
    add_styled_table(doc,
        ["Method", "Endpoint", "Auth", "Deskripsi"],
        [
            ["GET", "/panel/api/bps/health", "Yes", "Service health check"],
            ["GET/POST", "/panel/api/bps/config", "Yes", "Get/Update API configuration"],
            ["GET/POST", "/panel/api/bps/areas", "Yes", "List/Add monitored areas"],
            ["GET/POST", "/panel/api/bps/variables", "Yes", "List/Add variables"],
            ["DELETE", "/panel/api/bps/variables/:id", "Yes", "Remove variable"],
            ["GET", "/panel/api/bps/data", "Yes", "Query statistical data"],
            ["POST", "/panel/api/bps/sync/full", "Admin", "Full synchronization"],
            ["POST", "/panel/api/bps/sync/incremental", "Admin", "Incremental sync"],
            ["GET", "/panel/api/bps/sync/history", "Yes", "Sync operation history"],
            ["GET", "/panel/api/bps/catalog/search", "Yes", "Search BPS catalog"],
        ],
        col_widths=[0.7, 2.5, 0.6, 2.5]
    )

    # 8. Auth & Security
    section_heading(doc, "8. Authentication & Authorization")

    section_heading(doc, "8.1 JWT Authentication Flow", level=2)
    doc.add_paragraph(
        "1. User login dengan email + password\n"
        "2. Server validates credentials (bcrypt compare)\n"
        "3. Server generates JWT access token (expires: 1 hour)\n"
        "4. Server generates refresh token (expires: 30 days, httpOnly cookie)\n"
        "5. Client stores access token in memory\n"
        "6. Client sends Authorization: Bearer <token> pada setiap request\n"
        "7. Token expired \u2192 client calls /auth/refresh \u2192 new token pair\n"
        "8. Refresh token rotation: old token invalidated setelah refresh"
    )

    section_heading(doc, "8.2 RBAC Role Matrix", level=2)
    add_styled_table(doc,
        ["Role", "Dashboard", "Data CRUD", "BPS/Analytics", "Admin", "Verification", "Users/Perms"],
        [
            ["super_admin", "\u2713", "\u2713", "\u2713", "\u2713", "\u2713", "\u2713"],
            ["internal_admin", "\u2713", "\u2713", "\u2713", "\u2713", "\u2713", "\u2717"],
            ["pelaku_usaha", "\u2713", "Read", "Read", "\u2717", "\u2717", "\u2717"],
            ["pengolah_data", "\u2713", "\u2713", "\u2713", "\u2717", "\u2717", "\u2717"],
            ["internal_group", "\u2713", "Read", "Read", "\u2717", "\u2717", "\u2717"],
            ["guest", "\u2717", "\u2717", "\u2717", "\u2717", "\u2717", "\u2717"],
        ],
        col_widths=[1.2, 0.9, 0.9, 0.9, 0.9, 0.9, 0.9]
    )

    # 9. Security
    section_heading(doc, "9. Security Implementation")
    doc.add_paragraph(
        "Implementasi keamanan mengikuti standar OWASP Top 10 dan best practices DevSecOps."
    )

    add_styled_table(doc,
        ["Kategori", "Implementasi", "Detail"],
        [
            ["Authentication", "JWT + bcrypt", "Token expiration, refresh rotation, password hashing"],
            ["Authorization", "RBAC", "6-level roles, module-level permissions, field-level granularity"],
            ["Transport", "HTTPS/TLS", "Helmet HSTS, secure cookies, CSP headers"],
            ["Input Validation", "Zod + Sanitization", "Schema validation, XSS prevention, SQL injection prevention"],
            ["Rate Limiting", "express-rate-limit", "Login: 100/15min, configurable per endpoint"],
            ["CORS", "Whitelist", "Configurable origin list via environment"],
            ["File Upload", "Multer + Validation", "File type checking, size limits, storage isolation"],
            ["Audit", "Comprehensive Logging", "All actions logged, login attempts tracked, API call audit"],
            ["Dependencies", "ESLint Security", "Static analysis, dependency vulnerability scanning"],
        ],
        col_widths=[1.2, 1.5, 3.8]
    )

    # 10-14 more sections (abbreviated for length)
    section_heading(doc, "10. Integration Architecture")
    doc.add_paragraph(
        "Sistem menggunakan adapter pattern untuk integrasi dengan sistem eksternal. "
        "Setiap adapter mengimplementasikan base interface yang standar."
    )
    add_styled_table(doc,
        ["Sistem", "URL", "Auth Method", "Sync Type", "Status"],
        [
            ["BPS", "webapi.bps.go.id", "API Key", "Full & Incremental", "Active"],
            ["KOMINFO Tarif", "tariftel.komdigi.go.id", "API Key", "Periodic (cron)", "Active"],
            ["e-Telekomunikasi", "etelekomunikasi.komdigi.go.id", "Service Key", "Real-time", "Active"],
            ["OSS", "oss.go.id", "OAuth2", "On-demand", "Framework Ready"],
            ["POSTEL", "postel.go.id", "API Key", "Batch", "Framework Ready"],
            ["SDPPI", "sdppi.go.id", "API Key", "Batch", "Framework Ready"],
        ],
        col_widths=[1.2, 2.0, 1.0, 1.2, 1.0]
    )

    section_heading(doc, "11. Performance Requirements")
    add_styled_table(doc,
        ["Metrik", "Target", "Metode"],
        [
            ["API Response Time", "< 500ms (p95)", "Connection pooling, query optimization"],
            ["Page Load Time", "< 3 seconds", "Vite code splitting, lazy loading"],
            ["Concurrent Users", "100+ simultaneous", "PostgreSQL connection pool, stateless API"],
            ["Database Query", "< 200ms", "Indexed queries, materialized views"],
            ["WebSocket Latency", "< 100ms", "Event-driven architecture"],
            ["File Upload", "< 10MB per file", "Multer size limits, streaming"],
            ["Uptime", "99.5%", "PM2 process management, health checks"],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    section_heading(doc, "12. Deployment Architecture")
    doc.add_paragraph(
        "- Container: Docker (Dockerfile + Dockerfile.frontend)\n"
        "- Process Manager: PM2 (ecosystem.config.js)\n"
        "- Reverse Proxy: Apache/Nginx\n"
        "- Service Management: Systemd (telkom-insight-hub.service)\n"
        "- Frontend: Vite build output served at /panel/ path\n"
        "- Backend: Node.js Express server on port 4000\n"
        "- Database: PostgreSQL 15+ on port 5432\n"
        "- Static Files: storage/uploads/ for file uploads"
    )

    section_heading(doc, "13. Testing Strategy")
    doc.add_paragraph(
        "- Unit Testing: Vitest 3.2.4 dengan @testing-library/react\n"
        "- Component Testing: React Testing Library\n"
        "- Static Analysis: TypeScript strict mode + ESLint + Security plugin\n"
        "- Pre-commit: Husky 9.1.7 + lint-staged 16.1.4\n"
        "- UAT: 528 skenario test across 24 modul\n"
        "- Security: OWASP ZAP, dependency scanning"
    )

    section_heading(doc, "14. Monitoring & Logging")
    doc.add_paragraph(
        "- Audit Logging: Semua aksi pengguna tercatat di audit_logs\n"
        "- API Monitoring: Tracking request/response pada api_integration_logs\n"
        "- Login Tracking: Pencatatan login attempts untuk deteksi brute-force\n"
        "- Performance: Custom hooks (useAPIMonitoring, useMonitoring)\n"
        "- Health Check: GET /health endpoint untuk uptime monitoring\n"
        "- Security Dashboard: DevSecOps metrics visualization"
    )

    # Environment Variables
    section_heading(doc, "15. Environment Configuration")
    add_styled_table(doc,
        ["Variable", "Required", "Default", "Deskripsi"],
        [
            ["PORT", "Yes", "4000", "Server port"],
            ["NODE_ENV", "Yes", "development", "Environment mode"],
            ["DATABASE_URL", "Yes", "-", "PostgreSQL connection string"],
            ["JWT_SECRET", "Yes", "-", "JWT signing secret (32+ chars)"],
            ["ACCESS_EXPIRES_IN", "No", "1h", "Access token expiration"],
            ["REFRESH_EXPIRES_DAYS", "No", "30", "Refresh token expiration (days)"],
            ["CORS_ORIGIN", "Yes", "-", "Allowed CORS origins (comma-separated)"],
            ["ETELEKOM_API_URL", "No", "-", "e-Telekomunikasi API base URL"],
            ["ETELEKOM_SERVICE_KEY", "No", "-", "Service-to-service API key"],
            ["KOMINFO_TARIF_API_KEY", "No", "-", "KOMINFO tariff API key"],
            ["BPS_API_KEY", "No", "-", "BPS API key from webapi.bps.go.id"],
            ["VITE_API_BASE_URL", "Yes", "-", "Frontend API base URL"],
            ["VITE_MAPBOX_PUBLIC_TOKEN", "No", "-", "Mapbox GL access token"],
        ],
        col_widths=[2.0, 0.7, 1.0, 2.8]
    )

    add_signoff_box(doc, "Lembar Persetujuan — Technical Specification")

    path = os.path.join(OUTPUT_DIR, "TECHNICAL_SPECIFICATION.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: BLUEPRINT.docx
# ═══════════════════════════════════════════════════════════════════════════

def create_blueprint():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_cover_page(doc, "BLUEPRINT", "Cetak Biru Sistem Panel Manajemen Data Telekomunikasi & Analitik Strategis", "TIH-BP-2026-001")

    # TOC
    section_heading(doc, "Daftar Isi")
    toc_items = [
        "1. Executive Summary", "2. Visi & Misi Sistem", "3. Business Context",
        "4. System Blueprint Overview", "5. Module Architecture", "6. Data Architecture",
        "7. Integration Blueprint", "8. Security Architecture", "9. User Journey Maps",
        "10. Deployment Blueprint", "11. Scalability Strategy", "12. Risk Mitigation"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # 1. Executive Summary
    section_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "Telkom Insight Hub adalah platform manajemen data telekomunikasi dan analitik strategis "
        "yang dikembangkan sebagai panel pendukung untuk sistem e-Telekomunikasi milik "
        "Kementerian Komunikasi dan Digital Republik Indonesia. Platform ini menyediakan "
        "kemampuan pengelolaan data terpadu untuk 8 jenis layanan telekomunikasi, integrasi "
        "dengan data statistik BPS, analisis potensi pasar telekomunikasi berbasis wilayah, "
        "serta modul pendukung operasional seperti registrasi perusahaan, support ticketing, "
        "dan FAQ management."
    )
    doc.add_paragraph(
        "Cetak biru ini mendokumentasikan arsitektur menyeluruh dari sistem, mencakup "
        "desain modul, aliran data, strategi integrasi, keamanan, dan rencana deployment "
        "untuk memastikan sistem dapat beroperasi secara reliable, scalable, dan secure."
    )

    # 2. Visi & Misi
    section_heading(doc, "2. Visi & Misi Sistem")

    section_heading(doc, "2.1 Visi", level=2)
    doc.add_paragraph(
        "Menjadi platform terpadu untuk pengelolaan dan analisis data telekomunikasi nasional "
        "yang mendukung pengambilan keputusan strategis di lingkungan Kementerian Komunikasi dan Digital."
    )

    section_heading(doc, "2.2 Misi", level=2)
    doc.add_paragraph(
        "1. Menyediakan Single Source of Truth untuk data perizinan telekomunikasi\n"
        "2. Mengintegrasikan data multi-sumber (BPS, KOMINFO, OSS, POSTEL, SDPPI)\n"
        "3. Menyajikan analitik potensi pasar telekomunikasi berbasis wilayah\n"
        "4. Mengotomasi proses bisnis registrasi dan verifikasi perusahaan\n"
        "5. Memastikan keamanan dan kepatuhan data sesuai standar pemerintah"
    )

    # 3. Business Context
    section_heading(doc, "3. Business Context")

    section_heading(doc, "3.1 Stakeholder Map", level=2)
    add_styled_table(doc,
        ["Stakeholder", "Peran", "Kebutuhan Utama"],
        [
            ["Dirjen PPI", "Sponsor Eksekutif", "Dashboard strategis, laporan berkala"],
            ["Direktorat Telekomunikasi", "Pengguna Utama", "Manajemen data perizinan, analisis"],
            ["Pelaku Usaha", "Pengguna Eksternal", "Registrasi, tracking izin, informasi"],
            ["Tim IT Komdigi", "Pengelola Teknis", "Maintenance, monitoring, deployment"],
            ["BPS", "Penyedia Data", "Integrasi API data statistik"],
            ["KOMINFO/Komdigi", "Penyedia Data", "Sinkronisasi data tarif"],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    section_heading(doc, "3.2 Business Capability Map", level=2)
    add_styled_table(doc,
        ["Kapabilitas", "Deskripsi", "Modul Pendukung"],
        [
            ["Data Governance", "Pengelolaan siklus hidup data perizinan", "Data Management, Excel Import/Export"],
            ["Regulatory Compliance", "Kepatuhan terhadap regulasi telekomunikasi", "Company Verification, RBAC"],
            ["Market Intelligence", "Analisis potensi pasar per wilayah", "Telecom Potential, BPS Integration"],
            ["Stakeholder Service", "Pelayanan kepada pelaku usaha", "Registration, Support, FAQ"],
            ["Decision Support", "Dukungan pengambilan keputusan", "Dashboard, Visualization, Reports"],
            ["System Integration", "Interoperabilitas antar sistem", "e-Telekomunikasi, BPS, KOMINFO"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    # 4. System Blueprint
    section_heading(doc, "4. System Blueprint Overview")
    doc.add_paragraph(
        "Sistem terdiri dari 4 layer utama yang saling terintegrasi:"
    )

    add_styled_table(doc,
        ["Layer", "Komponen", "Teknologi", "Tanggung Jawab"],
        [
            ["Presentation", "React SPA, Mapbox Maps, Recharts", "React 18, TypeScript, TailwindCSS", "User interface, data visualization"],
            ["Application", "REST API, WebSocket, Background Jobs", "Express.js, JWT, node-cron", "Business logic, authentication, scheduling"],
            ["Data", "RDBMS, File Storage", "PostgreSQL, Multer", "Data persistence, file management"],
            ["Integration", "API Adapters, Sync Services", "Axios, Custom Services", "External system communication"],
        ],
        col_widths=[1.0, 2.0, 2.0, 1.5]
    )

    # 5. Module Architecture
    section_heading(doc, "5. Module Architecture")

    modules = [
        ("5.1", "Manajemen Data Telekomunikasi",
         "Modul inti untuk pengelolaan data perizinan 8 jenis layanan telekomunikasi.",
         [
            ["Data CRUD", "Create, Read, Update, Delete data perizinan", "High"],
            ["Excel Import", "Bulk import data dari spreadsheet dengan column mapping", "High"],
            ["Excel Export", "Export data ke format Excel", "Medium"],
            ["Search & Filter", "Pencarian multi-kriteria (company, status, location, date)", "High"],
            ["Pagination", "Server-side pagination untuk dataset besar", "High"],
            ["Detail View", "Tampilan detail per record dengan dokumen terkait", "Medium"],
            ["Location Filter", "Filter berdasarkan provinsi & kabupaten", "Medium"],
         ]),
        ("5.2", "Registrasi & Verifikasi Perusahaan",
         "Workflow end-to-end untuk registrasi perusahaan pelaku usaha telekomunikasi.",
         [
            ["Public Registration", "Form registrasi dengan upload dokumen", "High"],
            ["Document Upload", "NIB, NPWP, Akta Pendirian, KTP", "High"],
            ["PIC Management", "Person in Charge data per company", "Medium"],
            ["Verification Workflow", "submitted \u2192 in_review \u2192 verified/rejected", "High"],
            ["Admin Dashboard", "Panel admin untuk review & approve", "High"],
            ["Correction Notes", "Feedback system untuk dokumen yang perlu diperbaiki", "Medium"],
         ]),
        ("5.3", "Integrasi Data BPS",
         "Integrasi langsung dengan Badan Pusat Statistik untuk data statistik makroekonomi.",
         [
            ["API Configuration", "Konfigurasi API key, rate limit, base URL", "High"],
            ["Area Management", "Pemilihan provinsi & kabupaten yang dimonitor", "High"],
            ["Variable Management", "Pencarian & registrasi variabel statistik", "High"],
            ["Data Synchronization", "Full & incremental sync dengan tracking", "High"],
            ["Catalog Search", "Pencarian katalog variabel BPS", "Medium"],
            ["Data Visualization", "Line/bar charts untuk time-series data", "Medium"],
         ]),
        ("5.4", "Analisis Potensi Pasar Telekomunikasi",
         "Algoritma scoring multi-dimensi untuk mengukur potensi pasar per wilayah.",
         [
            ["Market Activity Score", "Scoring berdasarkan jumlah & diversitas lisensi", "High"],
            ["Untapped Opportunity", "Analisis potensi berdasarkan populasi & gap", "High"],
            ["BPS Demand Score", "Blending data statistik BPS dengan scoring V2", "Medium"],
            ["Interactive Map", "Choropleth map dengan Mapbox GL", "High"],
            ["Radar/Pie Charts", "Multi-dimensional visualization", "Medium"],
            ["Configuration", "Adjustable weights & parameters", "Medium"],
         ]),
        ("5.5", "Support & FAQ",
         "Modul pelayanan pengguna melalui ticketing dan knowledge base.",
         [
            ["Ticket Creation", "Pembuatan tiket support dari interface publik", "High"],
            ["Real-time Messaging", "WebSocket-based conversation per ticket", "High"],
            ["Admin Dashboard", "Panel admin untuk assignment & tracking", "High"],
            ["SLA Tracking", "Monitoring SLA metrics per ticket", "Medium"],
            ["FAQ Management", "CRUD FAQ dengan kategori & search", "Medium"],
         ]),
    ]

    for num, title, desc, features in modules:
        section_heading(doc, f"{num} {title}", level=2)
        doc.add_paragraph(desc)
        add_styled_table(doc,
            ["Fitur", "Deskripsi", "Prioritas"],
            features,
            col_widths=[1.5, 3.5, 0.8]
        )
        doc.add_paragraph("")

    # 6. Data Architecture
    section_heading(doc, "6. Data Architecture")
    doc.add_paragraph(
        "Data architecture mengikuti prinsip:\n"
        "- Single Source of Truth: PostgreSQL sebagai penyimpanan utama\n"
        "- Data Normalization: Third Normal Form (3NF) untuk core tables\n"
        "- Referential Integrity: Foreign key constraints antar tabel\n"
        "- Audit Trail: Semua perubahan tercatat dengan timestamp dan user\n"
        "- Soft Delete: Record penting tidak dihapus permanen"
    )

    # 7. Integration Blueprint
    section_heading(doc, "7. Integration Blueprint")
    doc.add_paragraph(
        "Integrasi menggunakan Adapter Pattern yang memungkinkan penambahan integrasi "
        "baru tanpa mengubah core system. Setiap adapter mengimplementasikan interface "
        "standar: connect(), fetch(), transform(), store()."
    )

    add_styled_table(doc,
        ["Sistem", "Jenis Integrasi", "Frekuensi", "Data yang Dipertukarkan"],
        [
            ["e-Telekomunikasi", "Bi-directional API", "Real-time", "Data perizinan, autentikasi"],
            ["BPS (webapi.bps.go.id)", "Inbound API", "Scheduled + On-demand", "Data statistik, variabel, area"],
            ["KOMINFO Tarif", "Inbound API", "Periodic (cron)", "Data tarif telekomunikasi"],
            ["OSS", "Inbound API", "On-demand", "Data registrasi berusaha"],
            ["POSTEL", "Inbound Batch", "Scheduled", "Data perizinan legacy"],
            ["SDPPI", "Inbound API", "Scheduled", "Data alokasi spektrum frekuensi"],
        ],
        col_widths=[1.3, 1.3, 1.5, 2.5]
    )

    # 8-12
    section_heading(doc, "8. Security Architecture")
    doc.add_paragraph(
        "Security architecture dibangun berlapis (Defense in Depth):\n\n"
        "Layer 1 — Network: Firewall, HTTPS/TLS, CORS whitelist\n"
        "Layer 2 — Application: JWT authentication, RBAC authorization, rate limiting\n"
        "Layer 3 — Data: Parameterized queries, input validation, encryption at rest\n"
        "Layer 4 — Monitoring: Audit logging, intrusion detection, security dashboard"
    )

    section_heading(doc, "9. User Journey Maps")

    section_heading(doc, "9.1 Pelaku Usaha — Registrasi", level=2)
    doc.add_paragraph(
        "1. Akses halaman /register\n"
        "2. Isi data perusahaan (nama, NIB, NPWP, alamat)\n"
        "3. Upload dokumen pendukung (Akta, KTP, dll.)\n"
        "4. Isi data Person in Charge (PIC)\n"
        "5. Submit registrasi\n"
        "6. Terima notifikasi status (email/dashboard)\n"
        "7. Jika perlu koreksi \u2192 perbaiki dan re-submit\n"
        "8. Setelah verified \u2192 akses penuh ke sistem"
    )

    section_heading(doc, "9.2 Admin — Verifikasi Perusahaan", level=2)
    doc.add_paragraph(
        "1. Login sebagai internal_admin/super_admin\n"
        "2. Akses Company Management dashboard\n"
        "3. Filter perusahaan berdasarkan status (submitted/in_review)\n"
        "4. Review dokumen dan data perusahaan\n"
        "5. Approve, reject, atau request correction\n"
        "6. Tambahkan catatan/feedback jika diperlukan\n"
        "7. Status otomatis terupdate di dashboard pelaku usaha"
    )

    section_heading(doc, "9.3 Pengolah Data — Analisis Pasar", level=2)
    doc.add_paragraph(
        "1. Login sebagai pengolah_data\n"
        "2. Akses Telecom Potential dashboard\n"
        "3. Pilih tahun dan parameter scoring\n"
        "4. View interactive map (choropleth by tier)\n"
        "5. Click area untuk detail breakdown\n"
        "6. Analisis radar chart multi-dimensi\n"
        "7. Export data ke Excel jika diperlukan"
    )

    section_heading(doc, "10. Deployment Blueprint")
    doc.add_paragraph(
        "Production Environment:\n"
        "- Server: Linux (Ubuntu 22.04 LTS recommended)\n"
        "- Container: Docker dengan multi-stage build\n"
        "- Process Manager: PM2 untuk auto-restart dan clustering\n"
        "- Reverse Proxy: Nginx/Apache untuk SSL termination\n"
        "- Database: PostgreSQL 15+ dengan regular backup\n"
        "- Monitoring: Health check endpoint + log aggregation\n"
        "- CI/CD: GitHub Actions / GitLab CI recommended"
    )

    section_heading(doc, "11. Scalability Strategy")
    doc.add_paragraph(
        "- Horizontal Scaling: PM2 cluster mode untuk multiple Node.js instances\n"
        "- Database: Connection pooling (pg pool), read replicas untuk query-heavy loads\n"
        "- Caching: React Query client-side caching, HTTP caching headers\n"
        "- CDN: Static assets served via CDN for frontend\n"
        "- Background Jobs: Queue-based processing untuk heavy tasks\n"
        "- API: Pagination dan lazy loading untuk large datasets"
    )

    section_heading(doc, "12. Risk Mitigation")
    add_styled_table(doc,
        ["Risiko", "Dampak", "Probabilitas", "Mitigasi"],
        [
            ["Database failure", "High", "Low", "Regular backup, failover setup"],
            ["API rate limit exceeded", "Medium", "Medium", "Rate limiting, queue-based sync"],
            ["Security breach", "High", "Low", "Defense in depth, audit logging, OWASP compliance"],
            ["External API downtime", "Medium", "Medium", "Retry mechanism, graceful degradation"],
            ["Data inconsistency", "High", "Low", "Transaction management, data validation"],
            ["Performance degradation", "Medium", "Medium", "Monitoring, indexing, caching"],
        ],
        col_widths=[1.5, 0.8, 0.8, 3.3]
    )

    add_signoff_box(doc, "Lembar Persetujuan — Blueprint")

    path = os.path.join(OUTPUT_DIR, "BLUEPRINT.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT 3: BUSINESS_REQUIREMENTS.docx
# ═══════════════════════════════════════════════════════════════════════════

def create_business_requirements():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_cover_page(doc, "BUSINESS\nREQUIREMENTS", "Dokumen Kebutuhan Bisnis — Sistem Panel Manajemen Data Telekomunikasi", "TIH-BRD-2026-001")

    # TOC
    section_heading(doc, "Daftar Isi")
    toc_items = [
        "1. Pendahuluan", "2. Business Objectives", "3. Scope & Boundaries",
        "4. Functional Requirements", "5. Non-Functional Requirements",
        "6. Business Rules", "7. User Stories & Acceptance Criteria",
        "8. Data Requirements", "9. Integration Requirements",
        "10. Compliance Requirements", "11. Training & Change Management",
        "12. Success Criteria"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # 1. Pendahuluan
    section_heading(doc, "1. Pendahuluan")

    section_heading(doc, "1.1 Latar Belakang", level=2)
    doc.add_paragraph(
        "Kementerian Komunikasi dan Digital melalui Direktorat Jenderal Penyelenggaraan Pos "
        "dan Informatika memerlukan sebuah sistem panel manajemen data yang terintegrasi "
        "untuk mendukung pengelolaan data perizinan telekomunikasi, analisis potensi pasar, "
        "dan pelayanan kepada pelaku usaha telekomunikasi. Saat ini, data tersebar di berbagai "
        "sistem (e-Telekomunikasi, POSTEL, OSS, SDPPI) tanpa platform terpadu untuk analisis "
        "dan pengambilan keputusan."
    )

    section_heading(doc, "1.2 Tujuan Dokumen", level=2)
    doc.add_paragraph(
        "Dokumen Business Requirements Document (BRD) ini bertujuan untuk:\n"
        "- Mendefinisikan kebutuhan bisnis yang harus dipenuhi oleh sistem\n"
        "- Menetapkan ruang lingkup dan batasan pengembangan\n"
        "- Mendokumentasikan aturan bisnis dan kriteria penerimaan\n"
        "- Menjadi dasar untuk pengembangan, testing, dan acceptance"
    )

    # 2. Business Objectives
    section_heading(doc, "2. Business Objectives")
    add_styled_table(doc,
        ["ID", "Objective", "KPI", "Target"],
        [
            ["BO-01", "Menyediakan Single Source of Truth data telekomunikasi", "Data completeness rate", "> 95%"],
            ["BO-02", "Mengintegrasikan data statistik BPS", "Variabel terhubung", "> 50 variabel"],
            ["BO-03", "Analisis potensi pasar per wilayah", "Cakupan wilayah", "34 provinsi"],
            ["BO-04", "Otomasi verifikasi perusahaan", "Waktu verifikasi", "< 3 hari kerja"],
            ["BO-05", "Pelayanan support real-time", "Response time", "< 4 jam"],
            ["BO-06", "Dashboard analitik real-time", "Data freshness", "< 1 jam"],
            ["BO-07", "Keamanan data sesuai standar", "Security incidents", "0 critical"],
            ["BO-08", "Integrasi seamless dengan e-Telekomunikasi", "Sync success rate", "> 99%"],
        ],
        col_widths=[0.6, 2.5, 1.5, 1.5]
    )

    # 3. Scope
    section_heading(doc, "3. Scope & Boundaries")

    section_heading(doc, "3.1 In Scope", level=2)
    doc.add_paragraph(
        "1. Manajemen data perizinan 8 jenis layanan telekomunikasi\n"
        "2. Registrasi dan verifikasi perusahaan pelaku usaha\n"
        "3. Integrasi data BPS untuk analisis statistik\n"
        "4. Sinkronisasi data tarif dari KOMINFO\n"
        "5. Analisis potensi pasar telekomunikasi berbasis wilayah\n"
        "6. Dashboard dan visualisasi data (charts, maps)\n"
        "7. Support ticketing dengan real-time messaging\n"
        "8. FAQ management\n"
        "9. User management dan RBAC\n"
        "10. Integrasi dengan sistem e-Telekomunikasi"
    )

    section_heading(doc, "3.2 Out of Scope", level=2)
    doc.add_paragraph(
        "1. Penerbitan izin (dilakukan di sistem e-Telekomunikasi)\n"
        "2. Pembayaran/billing (tidak ada fitur transaksi keuangan)\n"
        "3. Mobile application (scope saat ini: web-based)\n"
        "4. SMS/email notification service (belum termasuk)\n"
        "5. Public API untuk pihak ketiga (belum termasuk)"
    )

    # 4. Functional Requirements
    section_heading(doc, "4. Functional Requirements")

    fr_sections = [
        ("4.1", "Manajemen Data Telekomunikasi", [
            ["FR-DM-01", "Sistem harus mampu melakukan CRUD data perizinan untuk 8 jenis layanan", "Must Have"],
            ["FR-DM-02", "Sistem harus mendukung import data dari file Excel dengan column mapping", "Must Have"],
            ["FR-DM-03", "Sistem harus mendukung export data ke format Excel", "Must Have"],
            ["FR-DM-04", "Sistem harus menyediakan search multi-kriteria (company, status, location, date)", "Must Have"],
            ["FR-DM-05", "Sistem harus mendukung pagination server-side untuk dataset besar", "Must Have"],
            ["FR-DM-06", "Sistem harus menampilkan detail record dengan dokumen terkait", "Should Have"],
            ["FR-DM-07", "Sistem harus mendukung filter berdasarkan provinsi dan kabupaten", "Must Have"],
        ]),
        ("4.2", "Registrasi & Verifikasi Perusahaan", [
            ["FR-RV-01", "Sistem harus menyediakan form registrasi publik dengan upload dokumen", "Must Have"],
            ["FR-RV-02", "Sistem harus mendukung upload NIB, NPWP, Akta Pendirian, KTP", "Must Have"],
            ["FR-RV-03", "Sistem harus menyediakan workflow verifikasi (submitted > in_review > verified/rejected)", "Must Have"],
            ["FR-RV-04", "Sistem harus menyediakan dashboard admin untuk review dan approve perusahaan", "Must Have"],
            ["FR-RV-05", "Sistem harus mendukung catatan koreksi dan feedback ke pelaku usaha", "Should Have"],
            ["FR-RV-06", "Sistem harus mengelola data Person in Charge (PIC) per perusahaan", "Must Have"],
        ]),
        ("4.3", "Integrasi Data BPS", [
            ["FR-BPS-01", "Sistem harus mampu terhubung dengan webapi.bps.go.id", "Must Have"],
            ["FR-BPS-02", "Sistem harus mendukung konfigurasi API key dan rate limit", "Must Have"],
            ["FR-BPS-03", "Sistem harus mampu melakukan full dan incremental sync", "Must Have"],
            ["FR-BPS-04", "Sistem harus menyediakan pencarian katalog variabel BPS", "Should Have"],
            ["FR-BPS-05", "Sistem harus menampilkan visualisasi data time-series (line/bar chart)", "Must Have"],
            ["FR-BPS-06", "Sistem harus menyimpan history sinkronisasi dengan status", "Should Have"],
        ]),
        ("4.4", "Analisis Potensi Pasar", [
            ["FR-TP-01", "Sistem harus mampu menghitung Market Activity Score per wilayah", "Must Have"],
            ["FR-TP-02", "Sistem harus mampu menghitung Untapped Opportunity Score", "Must Have"],
            ["FR-TP-03", "Sistem harus menampilkan interactive map (choropleth) per tier", "Must Have"],
            ["FR-TP-04", "Sistem harus mendukung radar chart dan pie chart untuk analisis", "Should Have"],
            ["FR-TP-05", "Sistem harus menyediakan konfigurasi bobot scoring yang adjustable", "Should Have"],
            ["FR-TP-06", "Sistem harus mendukung BPS data blending untuk scoring V2", "Could Have"],
        ]),
        ("4.5", "Support & FAQ", [
            ["FR-SF-01", "Sistem harus menyediakan pembuatan tiket support dari interface publik", "Must Have"],
            ["FR-SF-02", "Sistem harus mendukung real-time messaging via WebSocket", "Must Have"],
            ["FR-SF-03", "Sistem harus menyediakan admin dashboard untuk assignment tiket", "Must Have"],
            ["FR-SF-04", "Sistem harus tracking SLA metrics per tiket", "Should Have"],
            ["FR-SF-05", "Sistem harus menyediakan CRUD FAQ dengan kategori", "Must Have"],
        ]),
        ("4.6", "User & Permission Management", [
            ["FR-UP-01", "Sistem harus mendukung 6 role (super_admin, internal_admin, pelaku_usaha, pengolah_data, internal_group, guest)", "Must Have"],
            ["FR-UP-02", "Sistem harus menyediakan module-level permissions (read/create/update/delete)", "Must Have"],
            ["FR-UP-03", "Sistem harus menyediakan admin interface untuk assignment permissions", "Must Have"],
            ["FR-UP-04", "Sistem harus mendukung field-level access control", "Could Have"],
        ]),
        ("4.7", "Integrasi Eksternal", [
            ["FR-IE-01", "Sistem harus terintegrasi dengan e-Telekomunikasi via service key", "Must Have"],
            ["FR-IE-02", "Sistem harus mampu sinkronisasi data tarif dari KOMINFO", "Must Have"],
            ["FR-IE-03", "Sistem harus menyediakan dashboard monitoring status integrasi", "Should Have"],
            ["FR-IE-04", "Sistem harus mendukung adapter pattern untuk integrasi baru", "Should Have"],
        ]),
    ]

    for num, title, reqs in fr_sections:
        section_heading(doc, f"{num} {title}", level=2)
        add_styled_table(doc,
            ["ID", "Requirement", "Priority"],
            reqs,
            col_widths=[1.0, 4.0, 1.0]
        )
        doc.add_paragraph("")

    # 5. Non-Functional Requirements
    section_heading(doc, "5. Non-Functional Requirements")
    add_styled_table(doc,
        ["ID", "Kategori", "Requirement", "Target"],
        [
            ["NFR-01", "Performance", "API response time (p95)", "< 500ms"],
            ["NFR-02", "Performance", "Page load time", "< 3 detik"],
            ["NFR-03", "Availability", "System uptime", "99.5%"],
            ["NFR-04", "Scalability", "Concurrent users", "100+"],
            ["NFR-05", "Security", "OWASP Top 10 compliance", "Full compliance"],
            ["NFR-06", "Security", "Data encryption in transit", "TLS 1.2+"],
            ["NFR-07", "Security", "Password policy", "bcrypt hash, min 8 chars"],
            ["NFR-08", "Usability", "Browser support", "Chrome, Firefox, Edge (latest 2)"],
            ["NFR-09", "Usability", "Responsive design", "Desktop + Tablet"],
            ["NFR-10", "Maintainability", "Code documentation", "JSDoc + inline comments"],
            ["NFR-11", "Reliability", "Data backup frequency", "Daily"],
            ["NFR-12", "Compliance", "Data retention policy", "Sesuai regulasi pemerintah"],
        ],
        col_widths=[0.7, 1.2, 2.3, 2.3]
    )

    # 6. Business Rules
    section_heading(doc, "6. Business Rules")
    add_styled_table(doc,
        ["ID", "Rule", "Deskripsi"],
        [
            ["BR-01", "Company Verification", "Perusahaan baru harus melalui verifikasi admin sebelum mendapat akses penuh"],
            ["BR-02", "Document Requirement", "Registrasi wajib menyertakan NIB, NPWP, Akta Pendirian, dan KTP PIC"],
            ["BR-03", "Role Assignment", "Hanya super_admin yang dapat mengubah role pengguna lain"],
            ["BR-04", "Data Ownership", "Pelaku usaha hanya dapat melihat data yang relevan dengan perusahaannya"],
            ["BR-05", "Audit Trail", "Semua perubahan data harus tercatat di audit log dengan timestamp dan user"],
            ["BR-06", "Token Expiration", "Access token berlaku 1 jam, refresh token berlaku 30 hari"],
            ["BR-07", "Rate Limiting", "Login dibatasi 100 percobaan per 15 menit per IP"],
            ["BR-08", "File Upload", "Ukuran file maksimal 10MB, format: PDF, JPG, PNG"],
            ["BR-09", "Data Sync", "Sinkronisasi BPS dan KOMINFO harus memiliki mekanisme retry"],
            ["BR-10", "Scoring Update", "Scoring potensi pasar harus di-recompute setelah data sync"],
        ],
        col_widths=[0.6, 1.5, 4.3]
    )

    # 7. User Stories
    section_heading(doc, "7. User Stories & Acceptance Criteria")

    stories = [
        ("US-01", "Sebagai pelaku usaha, saya ingin mendaftarkan perusahaan saya melalui form online",
         "Given halaman registrasi terbuka\nWhen saya mengisi semua field wajib dan upload dokumen\nThen perusahaan terdaftar dengan status 'submitted'\nAnd saya menerima konfirmasi registrasi berhasil"),
        ("US-02", "Sebagai admin, saya ingin memverifikasi perusahaan yang mendaftar",
         "Given ada perusahaan dengan status 'submitted'\nWhen saya review dokumen dan klik 'Verify'\nThen status berubah menjadi 'verified'\nAnd pelaku usaha mendapat akses penuh"),
        ("US-03", "Sebagai pengolah data, saya ingin import data dari Excel",
         "Given saya memiliki file Excel dengan data perizinan\nWhen saya upload file dan mapping kolom\nThen data berhasil diimport ke sistem\nAnd saya dapat melihat data baru di tabel"),
        ("US-04", "Sebagai analis, saya ingin melihat peta potensi pasar telekomunikasi",
         "Given data perizinan dan BPS sudah tersinkronisasi\nWhen saya akses halaman Telecom Potential\nThen saya melihat peta choropleth dengan warna per tier\nAnd saya dapat klik area untuk melihat detail scoring"),
        ("US-05", "Sebagai pengguna, saya ingin membuat tiket support",
         "Given saya sudah login ke sistem\nWhen saya mengisi form support dan klik Submit\nThen tiket terbuat dengan status 'open'\nAnd saya dapat melihat conversation thread"),
    ]

    for sid, story, criteria in stories:
        section_heading(doc, f"{sid}: {story}", level=3)
        doc.add_paragraph(f"Acceptance Criteria:\n{criteria}")
        doc.add_paragraph("")

    # 8-12
    section_heading(doc, "8. Data Requirements")
    doc.add_paragraph(
        "- Data perizinan telekomunikasi: 8 jenis layanan, ~40 fields per record\n"
        "- Data perusahaan: identitas, dokumen, PIC, status verifikasi\n"
        "- Data statistik BPS: time-series per area per variabel per tahun\n"
        "- Data tarif: sinkronisasi periodik dari KOMINFO\n"
        "- Data wilayah: 34 provinsi, 500+ kabupaten/kota\n"
        "- Audit trail: seluruh aksi pengguna dengan metadata"
    )

    section_heading(doc, "9. Integration Requirements")
    doc.add_paragraph(
        "- e-Telekomunikasi: Real-time bi-directional data exchange\n"
        "- BPS API: Scheduled sync dengan incremental capability\n"
        "- KOMINFO Tarif API: Periodic sync via cron job\n"
        "- OSS/POSTEL/SDPPI: Framework ready, adapter pattern implemented\n"
        "- Semua integrasi harus memiliki error handling dan retry mechanism"
    )

    section_heading(doc, "10. Compliance Requirements")
    doc.add_paragraph(
        "- Kepatuhan terhadap UU Telekomunikasi dan peraturan turunannya\n"
        "- Kepatuhan terhadap UU Perlindungan Data Pribadi (UU PDP)\n"
        "- Kepatuhan terhadap standar keamanan informasi pemerintah\n"
        "- Audit trail sesuai dengan ketentuan pengelolaan arsip elektronik\n"
        "- Data retention sesuai regulasi yang berlaku"
    )

    section_heading(doc, "11. Training & Change Management")
    doc.add_paragraph(
        "- Training admin: Manajemen pengguna, verifikasi perusahaan, ticketing\n"
        "- Training pengolah data: Import/export data, visualisasi, BPS config\n"
        "- Training pelaku usaha: Registrasi, tracking status, support\n"
        "- Dokumentasi teknis dan user guide\n"
        "- Knowledge transfer ke tim operasional"
    )

    section_heading(doc, "12. Success Criteria")
    add_styled_table(doc,
        ["Kriteria", "Metrik", "Target", "Measurement"],
        [
            ["Data Completeness", "% data termigrasi", "> 95%", "Query count vs source"],
            ["System Performance", "Response time p95", "< 500ms", "API monitoring"],
            ["User Adoption", "Active users/month", "> 50", "Login tracking"],
            ["Data Quality", "Error rate sync", "< 1%", "Sync history logs"],
            ["Security", "Critical vulnerabilities", "0", "Security scan"],
            ["Availability", "System uptime", "> 99.5%", "Health check monitoring"],
            ["UAT Pass Rate", "Test scenarios passed", "> 95%", "UAT execution report"],
        ],
        col_widths=[1.3, 1.5, 1.0, 2.5]
    )

    add_signoff_box(doc, "Lembar Persetujuan — Business Requirements")

    path = os.path.join(OUTPUT_DIR, "BUSINESS_REQUIREMENTS.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT 4: MONTHLY_PROGRESS_REPORT.docx
# ═══════════════════════════════════════════════════════════════════════════

def create_monthly_progress_report():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_cover_page(doc, "MONTHLY PROGRESS\nREPORT", "Laporan Kemajuan Bulanan — April 2026", "TIH-MPR-2026-04")

    # Executive Summary
    section_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "Laporan ini merangkum kemajuan pengembangan sistem Telkom Insight Hub untuk periode "
        "April 2026. Secara keseluruhan, pengembangan berjalan sesuai dengan jadwal yang "
        "direncanakan. Modul inti (Manajemen Data Telekomunikasi, Registrasi & Verifikasi) "
        "telah selesai 100%. Modul integrasi (BPS, KOMINFO) mencapai 85%, dan modul analitik "
        "(Potensi Pasar, Visualisasi) mencapai 70%. Fokus bulan ini adalah penyelesaian "
        "integrasi BPS dan optimasi scoring algorithm."
    )

    # Overall Progress
    section_heading(doc, "2. Overall Progress")
    add_styled_table(doc,
        ["Modul", "Target", "Aktual", "Status", "Catatan"],
        [
            ["Authentication & RBAC", "100%", "100%", "\u2713 Selesai", "JWT + 6 roles implemented"],
            ["Manajemen Data Telekomunikasi", "100%", "100%", "\u2713 Selesai", "8 jenis layanan, CRUD + Excel"],
            ["Registrasi & Verifikasi", "100%", "100%", "\u2713 Selesai", "Full workflow + admin dashboard"],
            ["Integrasi BPS", "90%", "85%", "\u25CB On Track", "Catalog search & viz remaining"],
            ["Sinkronisasi KOMINFO Tarif", "90%", "85%", "\u25CB On Track", "Periodic sync tested"],
            ["Potensi Pasar Telekomunikasi", "80%", "70%", "\u25CB On Track", "V2 scoring in progress"],
            ["Dashboard & Visualisasi", "80%", "75%", "\u25CB On Track", "Main dashboard complete"],
            ["Support & Ticketing", "100%", "100%", "\u2713 Selesai", "WebSocket real-time"],
            ["FAQ Management", "100%", "100%", "\u2713 Selesai", "CRUD + categories"],
            ["User Management", "100%", "100%", "\u2713 Selesai", "Admin panel complete"],
            ["Permission Management", "95%", "90%", "\u25CB On Track", "Field-level in progress"],
            ["Integrasi e-Telekomunikasi", "90%", "85%", "\u25CB On Track", "Service key auth active"],
            ["DevSecOps & Security", "90%", "90%", "\u2713 Selesai", "Helmet, rate limit, audit"],
            ["Documentation", "80%", "75%", "\u25CB On Track", "20+ docs, API ref complete"],
        ],
        col_widths=[1.8, 0.6, 0.6, 1.0, 2.3]
    )

    # Milestone Achievement
    section_heading(doc, "3. Milestone Achievement")
    add_styled_table(doc,
        ["Milestone", "Target Date", "Actual", "Status"],
        [
            ["Core Architecture Setup", "Jan 2026", "Jan 2026", "\u2713 Completed"],
            ["Authentication & RBAC", "Jan 2026", "Jan 2026", "\u2713 Completed"],
            ["Telekom Data CRUD", "Feb 2026", "Feb 2026", "\u2713 Completed"],
            ["Company Registration", "Feb 2026", "Feb 2026", "\u2713 Completed"],
            ["BPS Integration Phase 1", "Mar 2026", "Mar 2026", "\u2713 Completed"],
            ["KOMINFO Tariff Sync", "Mar 2026", "Mar 2026", "\u2713 Completed"],
            ["Telecom Potential V1", "Mar 2026", "Apr 2026", "\u25CB Delayed 2 weeks"],
            ["Dashboard & Visualization", "Apr 2026", "Apr 2026", "\u25CB In Progress"],
            ["BPS Integration Phase 2", "Apr 2026", "Apr 2026", "\u25CB In Progress"],
            ["Telecom Potential V2", "Apr 2026", "May 2026", "\u25CB Planned"],
            ["UAT Preparation", "May 2026", "-", "\u25CB Planned"],
            ["Production Deployment", "Jun 2026", "-", "\u25CB Planned"],
        ],
        col_widths=[2.0, 1.2, 1.2, 1.5]
    )

    # Key Achievements
    section_heading(doc, "4. Key Achievements This Month")
    doc.add_paragraph(
        "1. BPS Configuration Module — Implementasi lengkap modul konfigurasi BPS dengan "
        "API key management, area monitoring, dan variable registration\n\n"
        "2. Telecom Potential Dashboard — Implementasi dashboard potensi pasar dengan "
        "interactive Mapbox map, radar chart, dan sortable table\n\n"
        "3. Backend Integration Services — Implementasi BPS API Fetcher Service (23KB), "
        "Data Normalizer Service (17KB), dan Telecom Potential V2 Service (19KB)\n\n"
        "4. Data Migration — Migrasi data lokasi (provinsi & kabupaten) dan data perizinan "
        "dari sistem legacy\n\n"
        "5. Security Enhancement — Implementasi DevSecOps monitoring dashboard dengan "
        "audit logging dan security metrics"
    )

    # Technical Metrics
    section_heading(doc, "5. Technical Metrics")
    add_styled_table(doc,
        ["Metrik", "Bulan Lalu", "Bulan Ini", "Trend"],
        [
            ["Total Pages/Views", "18", "22", "\u2191 +4"],
            ["Total Components", "48", "55+", "\u2191 +7"],
            ["Total API Endpoints", "52", "60+", "\u2191 +8"],
            ["Total Database Tables", "35", "40+", "\u2191 +5"],
            ["Backend Services", "6", "12", "\u2191 +6"],
            ["Test Coverage", "45%", "55%", "\u2191 +10%"],
            ["Documentation Files", "15", "20+", "\u2191 +5"],
            ["Code Quality (ESLint)", "92%", "95%", "\u2191 +3%"],
        ],
        col_widths=[1.8, 1.2, 1.2, 1.2]
    )

    # Issues & Risks
    section_heading(doc, "6. Issues & Risks")
    add_styled_table(doc,
        ["ID", "Issue/Risk", "Severity", "Status", "Mitigation"],
        [
            ["ISS-01", "BPS API rate limit causing slow sync", "Medium", "Mitigated", "Implemented queuing & retry"],
            ["ISS-02", "Telecom Potential V2 complexity", "Low", "In Progress", "Phased implementation"],
            ["ISS-03", "Large dataset pagination performance", "Medium", "Resolved", "Server-side pagination + indexing"],
            ["RSK-01", "External API downtime during sync", "Medium", "Monitored", "Graceful degradation implemented"],
            ["RSK-02", "Data inconsistency between sources", "Low", "Monitored", "Validation checks in normalizer"],
        ],
        col_widths=[0.7, 2.0, 0.8, 0.8, 2.0]
    )

    # Next Month Plan
    section_heading(doc, "7. Plan for Next Month (May 2026)")
    add_styled_table(doc,
        ["Task", "Priority", "Estimated Effort", "Assignee"],
        [
            ["Complete Telecom Potential V2 scoring", "High", "5 days", "Backend Team"],
            ["BPS data visualization enhancement", "High", "3 days", "Frontend Team"],
            ["Field-level permission implementation", "Medium", "3 days", "Backend Team"],
            ["UAT scenario preparation (528 scenarios)", "High", "5 days", "QA Team"],
            ["Performance optimization & load testing", "Medium", "3 days", "DevOps Team"],
            ["Documentation completion & review", "Medium", "2 days", "All Teams"],
            ["Security audit & penetration testing", "High", "3 days", "Security Team"],
            ["Mobile responsive optimization", "Low", "2 days", "Frontend Team"],
        ],
        col_widths=[2.5, 0.8, 1.2, 1.5]
    )

    # Resource Allocation
    section_heading(doc, "8. Resource Allocation")
    add_styled_table(doc,
        ["Area", "Allocated", "Utilized", "Utilization %"],
        [
            ["Frontend Development", "2 developers", "2 developers", "100%"],
            ["Backend Development", "2 developers", "2 developers", "100%"],
            ["DevOps & Infrastructure", "1 engineer", "1 engineer", "80%"],
            ["QA & Testing", "1 tester", "1 tester", "60%"],
            ["Project Management", "1 PM", "1 PM", "100%"],
        ],
        col_widths=[2.0, 1.5, 1.5, 1.0]
    )

    # Budget
    section_heading(doc, "9. Budget Status")
    add_styled_table(doc,
        ["Kategori", "Budget", "Spent", "Remaining", "% Used"],
        [
            ["Development", "-", "-", "-", "-"],
            ["Infrastructure", "-", "-", "-", "-"],
            ["Licensing", "-", "-", "-", "-"],
            ["Testing", "-", "-", "-", "-"],
            ["Total", "-", "-", "-", "-"],
        ],
        col_widths=[1.5, 1.2, 1.2, 1.2, 0.8]
    )
    doc.add_paragraph("* Budget details to be filled by project management", style="Intense Quote")

    # Action Items
    section_heading(doc, "10. Action Items")
    add_styled_table(doc,
        ["#", "Action Item", "Owner", "Due Date", "Status"],
        [
            ["1", "Finalize BPS integration Phase 2", "Backend Lead", "15 May 2026", "Open"],
            ["2", "Complete Telecom Potential V2", "Backend Lead", "20 May 2026", "Open"],
            ["3", "Prepare UAT scenarios document", "QA Lead", "25 May 2026", "Open"],
            ["4", "Security audit report", "DevOps", "30 May 2026", "Open"],
            ["5", "Stakeholder demo & feedback", "PM", "30 May 2026", "Open"],
        ],
        col_widths=[0.4, 2.5, 1.2, 1.2, 0.8]
    )

    add_signoff_box(doc, "Lembar Persetujuan — Monthly Progress Report")

    path = os.path.join(OUTPUT_DIR, "MONTHLY_PROGRESS_REPORT.docx")
    doc.save(path)
    print(f"Saved: {path}")


# ── Run All ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    create_technical_specification()
    create_blueprint()
    create_business_requirements()
    create_monthly_progress_report()
    print("\nAll 4 DOCX files generated successfully!")
